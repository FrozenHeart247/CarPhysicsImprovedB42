from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Roadcraft_Dynamics_B42_Technical_Overview_RU.docx"
JAR = ROOT / "build" / "dist" / "roadcraft-agent.jar"
COMPATIBILITY = ROOT / "build" / "dist" / "compatibility.json"
SANDBOX_OPTIONS = (
    ROOT.parent
    / "Contents"
    / "mods"
    / "RoadcraftDynamicsB42"
    / "42"
    / "media"
    / "sandbox-options.txt"
)
RU_SANDBOX = SANDBOX_OPTIONS.parent / "lua" / "shared" / "Translate" / "RU" / "Sandbox.json"

# compact_reference_guide preset
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120
LIST_MARKER_DXA = 269
LIST_TEXT_DXA = 540
LIST_HANGING_DXA = 271

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GRID = "CDD5DF"
WHITE = "FFFFFF"
CAUTION = "7A5A00"
RISK = "9B1C1C"
POSITIVE = "1F3A5F"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(
                cell,
                CELL_TOP_BOTTOM_DXA,
                CELL_START_END_DXA,
                CELL_TOP_BOTTOM_DXA,
                CELL_START_END_DXA,
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_cell_text(cell, *, size=9.5, bold=False, color=INK, align=None):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], *, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        table.rows[0].cells[index].text = text
        shade_cell(table.rows[0].cells[index], LIGHT_BLUE)
        style_cell_text(table.rows[0].cells[index], size=font_size, bold=True, color=DARK_BLUE)
    repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for index, text in enumerate(row_data):
            row.cells[index].text = str(text)
            style_cell_text(row.cells[index], size=font_size)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title: str, body: str, *, kind="info"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    fill = CALLOUT
    accent = DARK_BLUE
    if kind == "caution":
        fill = "FFF8E8"
        accent = CAUTION
    elif kind == "risk":
        fill = "FDEEEE"
        accent = RISK
    elif kind == "positive":
        fill = LIGHT_BLUE
        accent = POSITIVE
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    for run in p2.runs:
        set_run_font(run, size=10, color=INK)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def paragraph_border_bottom(paragraph, color=BLUE, size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def create_numbering(doc: Document, fmt: str, text: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(LIST_TEXT_DXA))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(LIST_TEXT_DXA))
    ind.set(qn("w:hanging"), str(LIST_HANGING_DXA))
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.extend([start, num_fmt, lvl_text, suff, p_pr, r_pr])
    abstract.append(level)
    # OOXML requires every abstractNum before the first concrete num mapping.
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text: str, num_id: int, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=11, color=INK, bold=True)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "ROADCRAFT DYNAMICS  ·  BUILD 42"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    prefix = fp.add_run("Техническое описание  ·  ")
    set_run_font(prefix, size=9, color=MUTED)
    add_field(fp, "PAGE")

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = ffp.add_run("Независимая реализация · предварительная версия")
    set_run_font(label, size=9, color=MUTED)


def add_body_paragraph(doc, text: str, *, bold_lead: str | None = None, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, color=INK, bold=True)
        tail = p.add_run(text[len(bold_lead) :])
        set_run_font(tail, size=11, color=INK, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK, italic=italic)
    return p


def add_code_block(doc, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F8FA")
    cell.text = ""
    for index, line in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9, color=INK)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def read_compatibility():
    return json.loads(COMPATIBILITY.read_text(encoding="utf-8-sig"))


def read_settings_rows() -> list[list[str]]:
    text = SANDBOX_OPTIONS.read_text(encoding="utf-8-sig")
    translations = json.loads(RU_SANDBOX.read_text(encoding="utf-8-sig"))
    rows = []
    pattern = re.compile(r"option RoadcraftDynamics\.(\w+)\s*\{(.*?)\}", re.S)
    page_names = {
        "RCD_CorePage": "Основное",
        "RCD_DrivetrainPage": "Трансмиссия",
        "RCD_ResistancePage": "Сопротивление",
        "RCD_TractionPage": "Сцепление",
        "RCD_ImpactPage": "Столкновения",
        "RCD_SteeringPage": "Рулевое управление",
    }
    for match in pattern.finditer(text):
        key, body = match.groups()
        type_match = re.search(r"type\s*=\s*(\w+)", body)
        default_match = re.search(r"default\s*=\s*([^,\s]+)", body)
        min_match = re.search(r"min\s*=\s*([^,\s]+)", body)
        max_match = re.search(r"max\s*=\s*([^,\s]+)", body)
        page_match = re.search(r"page\s*=\s*(\w+)", body)
        display = translations.get(f"Sandbox_RCD_{key}", key)
        tooltip = translations.get(f"Sandbox_RCD_{key}_tooltip", "")
        value_type = (type_match.group(1) if type_match else "?").replace("boolean", "да/нет")
        default = default_match.group(1) if default_match else "—"
        limits = "—"
        if min_match and max_match:
            limits = f"{min_match.group(1)}…{max_match.group(1)}"
        page = page_names.get(page_match.group(1) if page_match else "", "Прочее")
        rows.append([page, f"{display}\n[{key}]", f"{value_type}; {default}; {limits}", tooltip])
    return rows


def add_cover(doc: Document, compatibility):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ТЕХНИЧЕСКОЕ ОПИСАНИЕ")
    set_run_font(r, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("Roadcraft Dynamics")
    set_run_font(r, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    r = subtitle.add_run("Независимая физика транспорта для Project Zomboid Build 42")
    set_run_font(r, size=15, color=DARK_BLUE)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(28)
    r = note.add_run("Архитектура, механики, параметры, установка, совместимость и план проверки")
    set_run_font(r, size=10.5, color=MUTED, italic=True)

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(24)
    paragraph_border_bottom(line, BLUE, "14", "8")

    metadata = [
        ["Ветка игры", "Build 42 (общая папка 42)"],
        ["Проверенный снимок", f"{compatibility['knownTested']['gameVersion']} · {compatibility['knownTested']['gameBuild']}"],
        ["Состояние", "0.1.0-dev · статически проверено · игровой runtime-тест ещё не выполнен"],
        ["Дата документа", "31 августа 2026"],
    ]
    add_table(doc, ["Параметр", "Значение"], metadata, [2700, 6660], font_size=10)

    add_callout(
        doc,
        "Главная граница доказанного",
        "Сборка, математические тесты, контракт трансформера, реальная загрузка -javaagent без запуска игры и тест установщика прошли. Управляемость, Bullet-физика, сохранения и мультиплеер должны быть подтверждены внутри игры.",
        kind="caution",
    )
    doc.add_page_break()


def add_contents(doc: Document, decimal_id: int):
    doc.add_heading("Содержание", level=1)
    sections = [
        "Назначение и границы независимой реализации",
        "Функциональные механики",
        "Архитектура новой версии",
        "Подробная модель движения",
        "Настройки сервера и клиента",
        "Совместимость с обновлениями B42",
        "Установка, обновление и удаление",
        "Мультиплеер и владение физикой",
        "Производительность",
        "Проверки и статус готовности",
        "Ограничения, исключения и дальнейшая работа",
        "Происхождение и лицензионная граница",
    ]
    for item in sections:
        add_list_item(doc, item, decimal_id)
    add_callout(
        doc,
        "Как читать статусы",
        "«Реализовано» означает наличие новой логики и автоматических проверок. Это не равно «подтверждено в игре». Для последнего используется отдельная отметка runtime/SP/MP.",
    )


def add_scope(doc: Document, bullet_id: int):
    doc.add_heading("1. Назначение и границы независимой реализации", level=1)
    add_body_paragraph(
        doc,
        "Цель проекта — восстановить набор реалистичных механик транспорта для Build 42, но не возвращать удалённые бинарные классы исходного мода. Новая версия имеет другое имя, собственную структуру, собственное математическое ядро, новые тексты, настройки и установщик.",
    )
    doc.add_heading("1.1. Что считалось запросом пользователя", level=2)
    for text in [
        "Воспроизвести наблюдаемое поведение мода с нуля, без копирования исходных Java/Lua-файлов, ресурсов и бинарных классов.",
        "Сделать одну общую ветку Build 42, не привязанную жёстко к 42.20.4.",
        "Сохранить ручную установку Java-компонента, если без неё нельзя получить низкоуровневую физику.",
        "Подготовить подробный DOCX, чтобы было понятно, что мод меняет и как его проверять.",
    ]:
        add_list_item(doc, text, bullet_id)

    doc.add_heading("1.2. Что не считалось запросом", level=2)
    add_body_paragraph(
        doc,
        "Текстовый файл старого мода с указанием скопировать папку zombie в каталог игры рассматривался только как артефакт референса. Эта инструкция не выполнялась и не переносилась в новую схему установки.",
    )
    add_callout(
        doc,
        "Правовая оговорка",
        "Независимая реализация и отказ от копирования снижают риск, но этот документ не является юридическим заключением и не гарантирует отсутствие претензий. Перед публичным релизом полезно отдельно проверить название, описание, визуальные материалы и юрисдикцию распространения.",
        kind="caution",
    )


def add_functional_scope(doc: Document, bullet_id: int):
    doc.add_heading("2. Функциональные механики", level=1)
    add_body_paragraph(
        doc,
        "Ниже зафиксирован функциональный объём первой новой версии. Это требования по поведению, а не описание внутреннего кода старого мода.",
    )
    rows = [
        ["Двигатель и обороты", "Расчёт холостого хода, кривой момента, красной зоны и связи оборотов со скоростью колёс.", "Реализовано; требует калибровки в игре"],
        ["Автоматическая КПП", "Передачи вперёд, нейтраль, задний ход, повышение/понижение по оборотам.", "Реализовано"],
        ["Ручная КПП", "Клиентский режим и отдельные клавиши повышения/понижения; сервер может запретить.", "Реализовано"],
        ["Гидротрансформатор", "Усиление момента на старте и постепенное уменьшение проскальзывания.", "Реализовано"],
        ["Задний ход", "Мягкое ограничение скорости и безопасная смена направления около остановки.", "Реализовано"],
        ["Сцепление", "Износ/состояние шин, мокрая дорога, снег, бездорожье, ограничение тяги и торможения.", "Реализовано; визуальный эффект требует runtime"],
        ["Пробуксовка", "Burnout при превышении лимита тяги и блокировка задних колёс ручником.", "Реализовано; требует runtime"],
        ["Рулевое управление", "Плавная скорость поворота и возврата руля, меньшая резкость на скорости.", "Реализовано; требует настройки ощущений"],
        ["Сопротивления", "Аэродинамическое сопротивление, качение и повышенное сопротивление вне дороги.", "Реализовано"],
        ["Буксировка", "Ограниченная помощь на малой скорости и уменьшение нежелательного торможения прицепа.", "Реализовано; требует тестов составов"],
        ["Столкновения", "Отдельные масштабы импульса для растений, зомби и трупов; нормализация нативной массы.", "Реализовано; требует Bullet/runtime"],
        ["MP-authority", "Расчёт сил только у владельца локальной физической симуляции машины.", "Реализовано; MP не подтверждён"],
    ]
    add_table(doc, ["Подсистема", "Что меняется", "Статус"], rows, [2100, 4800, 2460], font_size=9.1)

    doc.add_heading("2.1. Осознанно не включено в первую версию", level=2)
    for text in [
        "Таблица мощности и массы для конкретных моделей машин: новая логика выводит параметры из актуального VehicleScript и общих категорий.",
        "Заменяющие звуки двигателя, изображения и другие ресурсы референса.",
        "Изменение вместимости багажников или полная замена ItemContainer.",
        "Частные обходы ошибок для трупов животных и интеграции с конкретными сторонними модами.",
        "Копирование старых названий, строк интерфейса, настроечных констант и бинарных классов.",
    ]:
        add_list_item(doc, text, bullet_id)


def add_architecture(doc: Document, bullet_id: int):
    doc.add_heading("3. Архитектура новой версии", level=1)
    add_callout(
        doc,
        "Ключевое отличие",
        "В каталоге игры не размещаются классы zombie/**. Java-agent добавляет узкие вызовы в загружаемые классы в памяти, а при несовместимой структуре оставляет ванильную физику.",
        kind="positive",
    )
    rows = [
        ["Lua-мод", "Общая папка 42, sandbox-настройки, Mod Options, клавиши ручной КПП, передача конфигурации в Java."],
        ["Java-agent", "Запускается через -javaagent, один раз анализирует целевые классы при их загрузке и внедряет только заранее проверенные точки вызова."],
        ["Transformer", "Проверяет сигнатуры методов и нужные вызовы Bullet; при несовпадении не заменяет класс целиком."],
        ["RoadcraftHooks", "Связывает текущий CarController/BaseVehicle с математической моделью, применяет силы и импульсы только при разрешённом authority."],
        ["PzAccess", "Версионированный адаптер к B42.20.4. Reflection-метаданные кешируются, финальный Bullet dispatch использует точный MethodHandle; после трёх ошибок runtime отключается."],
        ["DrivetrainModel", "Чистая детерминированная Java-математика без импортов Project Zomboid; тестируется отдельно от игры."],
    ]
    add_table(doc, ["Слой", "Ответственность"], rows, [2300, 7060], font_size=9.5)

    doc.add_heading("3.1. Поток одного обновления машины", level=2)
    steps = [
        "Игра входит в CarController.update/updateTrailer; агент сохраняет контекст контроллера.",
        "Ванильный код читает управление и рассчитывает исходные engineForce, brakingForce и steering.",
        "Перед Bullet.controlVehicle вызов перенаправляется в RoadcraftHooks.",
        "Адаптер проверяет готовность конфигурации, локальное владение физикой, состояние двигателя, колёс, покрытия и погоды.",
        "DrivetrainModel возвращает тягу, торможение, передачу, обороты, пробуксовку, сопротивления и допустимый угол руля.",
        "RoadcraftHooks переводит результат в игровые единицы и вызывает нативный Bullet; дополнительное сопротивление прикладывается отдельной силой.",
        "После выхода временный контекст очищается; состояние конкретной машины сохраняется в слабой карте, не удерживающей удалённые машины.",
    ]
    decimal_id = create_numbering(doc, "decimal", "%1.")
    for step in steps:
        add_list_item(doc, step, decimal_id)

    doc.add_heading("3.2. Узкие точки внедрения", level=2)
    hooks = [
        ["CarController.update() / updateTrailer()", "Контекст + замена конкретного вызова Bullet.controlVehicle(int,float,float,float)."],
        ["WorldSimulation.updateVehiclePhysics()", "Подготовка общей нормализации физической массы до шага Bullet."],
        ["BaseVehicle: импульсы растений/персонажей/трупов", "Локальное масштабирование только записей, появившихся в выбранном scope."],
        ["LuaManager.init()", "Регистрация префиксных глобальных функций RCD_* для безопасного Lua-моста."],
    ]
    add_table(doc, ["Цель", "Что именно добавляется"], hooks, [3400, 5960], font_size=9.5)


def add_physics_details(doc: Document, bullet_id: int):
    doc.add_heading("4. Подробная модель движения", level=1)
    sections = [
        (
            "4.1. Двигатель, передачи и гидротрансформатор",
            [
                "Момент двигателя задаётся новой кусочно-сглаженной кривой: от холостого хода к пику и далее к красной зоне с мягким ограничителем.",
                "Расчётные обороты стремятся к оборотам, связанным с радиусом колеса, текущей передачей и главной парой; на автомате допускается проскальзывание гидротрансформатора.",
                "Автомат повышает и понижает передачу по порогам RPM, но не меняет направление на заметной скорости. Ручной режим хранит запрос отдельно для ID машины.",
                "При заглушенном двигателе AutoStart может запросить запуск исправного двигателя при газе; при отключённом AutoStart новая тяга не запускает мотор самопроизвольно.",
                "Ограничение заднего хода снижает тягу в приближении к заданной скорости, а не мгновенно обнуляет скорость.",
            ],
        ),
        (
            "4.2. Сцепление, шины и погода",
            [
                "Доступная продольная сила ограничена эффективным коэффициентом сцепления, массой и долей нагрузки на ведущие колёса.",
                "Коэффициент формируется из общего множителя, состояния установленных шин, мокроты, силы снега и доли бездорожья.",
                "Количество колёс берётся динамически из VehicleScript. Задние колёса определяются по признаку front, а не по жёстким индексам.",
                "Если требуемая тяга выше предела, накапливается wheelSlip; выходная тяга уменьшается. При восстановлении сцепления slip спадает плавно.",
                "Ручник и задняя часть рабочего тормоза сравниваются с доступным сцеплением задней оси; превышение отмечается как блокировка задних колёс.",
            ],
        ),
        (
            "4.3. Рулевое управление",
            [
                "На малой скорости колёса реагируют быстрее, на высокой — скорость изменения и максимальная резкость уменьшаются.",
                "Поворот, возврат к центру и смена направления имеют отдельные коэффициенты. Между 0 и SteeringHighSpeedReference применяется плавная интерполяция.",
                "Пробуксовка дополнительно уменьшает эффективный угол руля, чтобы не создавать нереалистичное мгновенное сцепление.",
            ],
        ),
        (
            "4.4. Сопротивление и торможение",
            [
                "Аэродинамическая сила растёт пропорционально квадрату скорости и зависит от категории машины.",
                "Сопротивление качению содержит постоянную и скоростную составляющие; на бездорожье используются повышенные значения.",
                "Рабочий и стояночный тормоза ограничиваются общим доступным сцеплением, чтобы сила торможения не была физически бесконечной.",
            ],
        ),
        (
            "4.5. Столкновения, масса и буксировка",
            [
                "Импульсы от растений, живых персонажей/зомби и трупов масштабируются отдельно и только внутри соответствующего вызова BaseVehicle.",
                "Нативная масса нормализуется детерминированным коэффициентом от DynamicMassReference; она не повышается выше массы VehicleScript.",
                "EasyTow уменьшает нежелательное торможение контроллера прицепа и добавляет ограниченную помощь на малой скорости, но не отключает массу или столкновения.",
                "Глобальные VehicleScript не мутируются, поэтому одна локальная машина не меняет параметры всех экземпляров того же scriptName.",
            ],
        ),
    ]
    for title, bullets in sections:
        doc.add_heading(title, level=2)
        for text in bullets:
            add_list_item(doc, text, bullet_id)


def add_settings(doc: Document):
    doc.add_page_break()
    doc.add_heading("5. Настройки сервера и клиента", level=1)
    add_body_paragraph(
        doc,
        "Sandbox-параметры являются серверной/мировой конфигурацией. Mod Options управляют только личным режимом коробки, автоматическим выбором направления и клавишами. Серверный запрет ручной КПП имеет приоритет над клиентом.",
    )
    client_rows = [
        ["Ручная коробка передач", "Включает ручной режим для локального игрока, если сервер разрешил."],
        ["Автоматически включать задний ход", "После остановки торможение может выбрать R, а газ — передачу вперёд."],
        ["Повысить/понизить передачу", "Переназначаемые клавиши; по умолчанию стрелки вверх/вниз."],
    ]
    add_table(doc, ["Клиентская опция", "Назначение"], client_rows, [3100, 6260], font_size=9.5)

    doc.add_heading("5.1. Полный перечень Sandbox-настроек", level=2)
    add_body_paragraph(
        doc,
        "В колонке «тип; default; диапазон» указаны значения из текущего sandbox-options.txt. Нули разрешены у некоторых множителей как сознательный способ отключить соответствующий эффект.",
    )
    rows = read_settings_rows()
    add_table(
        doc,
        ["Раздел", "Параметр [ключ]", "Тип; default; диапазон", "Что делает"],
        rows,
        [1400, 2700, 1850, 3410],
        font_size=8.2,
    )


def add_compatibility(doc: Document, compatibility, bullet_id: int):
    doc.add_heading("6. Совместимость с обновлениями B42", level=1)
    add_body_paragraph(
        doc,
        "Папка мода называется 42 и mod.info задаёт versionMin=42.0. Это общая ветка B42, но не обещание, что любой будущий патч автоматически совместим.",
    )
    rows = [
        ["Точный проверенный снимок", "Хэши целевых классов совпали", "ACTIVE", "Полная активация после получения Lua-конфигурации."],
        ["Другая B42, структура прежняя", "Методы, сигнатуры и вызовы совпали, но хэш класса новый", "ACTIVE_UNVERIFIED", "Мод работает в режиме непроверенной совместимости; нужен игровой smoke-test."],
        ["Структура изменилась", "Нет метода, другая сигнатура или исчез нужный callsite", "INCOMPATIBLE", "Трансформация отклоняется; сохраняется ванильное поведение."],
        ["Java-agent отсутствует", "Lua не видит RCD_* bridge", "NOT_INSTALLED", "Мод сообщает об этом в console.txt, физика ванильная."],
        ["Lua ещё не передал настройки", "Хуки загружены, конфигурация не активирована", "WAITING_CONFIG", "Новая физика не применяется."],
    ]
    add_table(doc, ["Сценарий", "Проверка", "Статус", "Результат"], rows, [1850, 2410, 2200, 2900], font_size=8.8)

    doc.add_heading("6.1. Что делать после обновления игры", level=2)
    for text in [
        "Не менять номер Java classfile и не подменять хэши вручную: это не исправляет изменившийся API.",
        "Проверить, сохранилась ли строка -javaagent в launcher после обновления игры, и при необходимости восстановить её из backup.",
        "Запустить игру, найти в console.txt строку RoadcraftDynamics со статусом ACTIVE, ACTIVE_UNVERIFIED или INCOMPATIBLE.",
        "Для ACTIVE_UNVERIFIED пройти короткий SP-smoke: запуск, разгон, торможение, задний ход, бездорожье, столкновение и буксировка.",
        "При INCOMPATIBLE обновить только версионированный адаптер/трансформер под новую структуру B42; копировать новые классы zombie целиком не требуется.",
    ]:
        add_list_item(doc, text, bullet_id)

    tested = compatibility["knownTested"]
    class_rows = []
    for value in tested["classes"].values():
        class_rows.append([value["internalName"], value["sha256"]])
    doc.add_heading("6.2. Проверенные отпечатки 42.20.4", level=2)
    add_table(doc, ["Класс", "SHA-256"], class_rows, [4000, 5360], font_size=8.2)


def add_installation(doc: Document, decimal_id: int, bullet_id: int):
    doc.add_heading("7. Установка, обновление и удаление", level=1)
    add_callout(
        doc,
        "Workshop + ручная установка",
        "Lua-часть загружается непосредственно из Workshop item. Steam не может подключить Java-agent к launcher, поэтому только JAR и launcher устанавливаются вручную.",
        kind="positive",
    )
    doc.add_heading("7.1. Структура Workshop item", level=2)
    install_decimal_id = create_numbering(doc, "decimal", "%1.")
    steps = [
        "Workshop-папка содержит workshop.txt и Contents\\mods\\RoadcraftDynamicsB42.",
        "Версионированная Lua-часть находится в 42\\mod.info и 42\\media; она запускается непосредственно из Workshop item.",
        "Java-agent, compatibility metadata, инструкция, лицензия и SHA-256 находятся рядом в MANUAL_INSTALLATION.",
        "Папка Dev не входит в Contents и не отправляется подписчикам.",
    ]
    for step in steps:
        add_list_item(doc, step, install_decimal_id)
    add_code_block(
        doc,
        [
            "Contents\\mods\\RoadcraftDynamicsB42\\42\\mod.info",
            "Contents\\mods\\RoadcraftDynamicsB42\\42\\media\\...",
            "Contents\\mods\\RoadcraftDynamicsB42\\MANUAL_INSTALLATION\\roadcraft-agent.jar",
        ],
    )

    add_callout(
        doc,
        "Без пользовательских скриптов",
        "В Workshop payload нет install, verify или uninstall-скриптов. Пользователь выполняет короткую ручную установку и всегда видит, какие два файла и одна launcher-строка добавляются.",
        kind="positive",
    )

    doc.add_heading("7.2. Установка Java-agent", level=2)
    manual_decimal_id = create_numbering(doc, "decimal", "%1.")
    manual_steps = [
        "Закрыть игру и сделать backup реально используемого launcher-файла: ProjectZomboid64.json для Steam-клиента либо соответствующего client/server BAT.",
        "Через Steam открыть Управление → Просмотреть локальные файлы. Открывшаяся папка с projectzomboid.jar — это GameDirectory.",
        "Скопировать roadcraft-agent.jar и compatibility.json из MANUAL_INSTALLATION в <GameDirectory>\\RoadcraftDynamics.",
        "Для Steam-клиента добавить строку -javaagent:RoadcraftDynamics/roadcraft-agent.jar ровно один раз в верхний массив vmArgs файла ProjectZomboid64.json.",
        "Для BAT-запуска добавить marker-блок, сохраняющий существующий _JAVA_OPTIONS, перед первым запуском java.exe. Редактировать только фактически используемый launcher.",
        "Включить RoadcraftDynamicsB42 в Mod Manager или Mods= server-конфигурации; установить агент на каждый MP-клиент.",
        "Проверить свежий console.txt: ожидается ACTIVE или ACTIVE_UNVERIFIED. Точные JSON/BAT фрагменты находятся в MANUAL_INSTALLATION\\INSTALL_RU.txt.",
    ]
    for step in manual_steps:
        add_list_item(doc, step, manual_decimal_id)

    add_code_block(
        doc,
        [
            "<GameDirectory>\\RoadcraftDynamics\\roadcraft-agent.jar",
            '"-javaagent:RoadcraftDynamics/roadcraft-agent.jar"',
        ],
    )

    doc.add_heading("7.3. Обновление", level=2)
    add_body_paragraph(
        doc,
        "Steam обновляет Workshop item, но не может заменить JAR, вручную скопированный в каталог игры. После обновления мода нужно повторно скопировать roadcraft-agent.jar и compatibility.json из MANUAL_INSTALLATION. Launcher-строка остаётся прежней. После патча игры следует проверить, что launcher не был заменён, и затем проверить runtime-статус в console.txt.",
    )

    doc.add_heading("7.4. Удаление", level=2)
    add_body_paragraph(
        doc,
        "Закрыть игру, удалить только собственную строку -javaagent или точный marker-блок, затем удалить <GameDirectory>\\RoadcraftDynamics. Lua-часть удаляется обычной отпиской от Workshop item. projectzomboid.jar не изменяется.",
    )


def add_multiplayer(doc: Document, bullet_id: int):
    doc.add_heading("8. Мультиплеер и владение физикой", level=1)
    add_body_paragraph(
        doc,
        "Низкоуровневые силы нельзя безусловно рассчитывать и на сервере, и на каждом клиенте. Перед применением новой тяги адаптер вызывает BaseVehicle.isLocalPhysicSim(). Только узел, которому игра отдала локальную физическую симуляцию этой машины, выполняет расчёт.",
    )
    rows = [
        ["Одиночная игра", "Локальный процесс", "Один расчёт сил."],
        ["Listen-host", "Host или назначенный локальный владелец", "Не применять повторно на удалённом клиенте."],
        ["Dedicated server", "Server-authority для серверных машин; локальный client authority там, где его выдаёт игра", "Нужна отдельная установка -javaagent в server launch script."],
        ["Пассажир/наблюдатель", "Не владеет физикой", "Мод не прикладывает силы."],
    ]
    add_table(doc, ["Режим", "Кто считает", "Ожидаемое поведение"], rows, [2200, 3400, 3760], font_size=9.2)
    doc.add_heading("8.1. Что обязательно проверить", level=2)
    for text in [
        "Одинаковая конфигурация sandbox на сервере и передача её клиентам.",
        "Отсутствие двойного ускорения/торможения при пересадке водителя и смене authority.",
        "Поведение host и удалённого клиента при столкновении с зомби, трупом и растением.",
        "Буксировка между чанками, выход/вход водителя и повторное подключение.",
        "Dedicated server запускается именно через batch/launcher, куда добавлен -javaagent.",
    ]:
        add_list_item(doc, text, bullet_id)


def add_performance(doc: Document, bullet_id: int):
    doc.add_heading("9. Производительность", level=1)
    add_body_paragraph(
        doc,
        "Java-agent не сканирует игру каждый кадр. Трансформация целевых классов происходит один раз при загрузке. В рабочем цикле остаётся дополнительный статический вызов, сбор состояния машины, чистый расчёт и применение результата к Bullet.",
    )
    rows = [
        ["Папка zombie с полной заменой", "Нет дополнительного dispatch к hook, но весь ванильный класс заменён. Цена несовместимости и сопровождения значительно выше."],
        ["Узкий Java-agent", "Один hook-dispatch и адаптер поверх ванильного класса. Сам механизм агента почти не влияет на steady-state, но текущий reflection-адаптер тяжелее полностью типизированной прямой реализации."],
        ["Общее", "Основную цену определяют сбор состояния, allocations, вызовы адаптера и native Bullet, а не одноразовая трансформация class-кода."],
    ]
    add_table(doc, ["Подход", "Практический эффект"], rows, [2600, 6760], font_size=9.5)

    doc.add_heading("9.1. Текущий микротест чистого ядра", level=2)
    perf_rows = [
        ["1", "10 000", "367,8 нс", "3,68 мс"],
        ["10", "100 000", "315,7 нс", "31,57 мс"],
        ["50", "500 000", "103,6 нс", "51,80 мс"],
    ]
    add_table(doc, ["Машин", "Обновлений", "Среднее на update", "Общее"], perf_rows, [1700, 2100, 2850, 2710], font_size=9.5)
    add_callout(
        doc,
        "Не считать это FPS-доказательством",
        "Микротест измеряет только DrivetrainModel после прогрева JVM. Он не включает reflection-адаптер, получение погоды/колёс, Bullet, сетевую репликацию, рендер и сборщик мусора игры. Снижение ns/update в больших циклах отражает JIT/форму теста, а не доказанное масштабирование.",
        kind="caution",
    )
    doc.add_heading("9.2. Известные расходы hot path", level=2)
    for text in [
        "Reflection Field/Method metadata создаётся один раз и кешируется, но Field.get/Method.invoke для состояния игры остаются в рабочем пути. Статическая оценка даёт несколько десятков вызовов для типичной четырёхколёсной машины; точное число зависит от ветки и требует профайлера.",
        "Финальный Bullet.controlVehicle больше не использует Method.invoke: после проверки сигнатуры он вызывается через кешированный MethodHandle.invokeExact.",
        "На обновление создаются снимки VehicleInput/VehicleState/VehicleOutput и при изменении конфигурации — PhysicsSettings; это кандидаты на оптимизацию после профилирования.",
        "Состояние хранится по объекту машины, а ThreadLocal используется только для краткого controller/impulse scope.",
        "Расчёт пропускается, если агент/конфигурация не готовы, мод отключён или машина не принадлежит локальной физической симуляции.",
    ]:
        add_list_item(doc, text, bullet_id)
    add_body_paragraph(
        doc,
        "Первый практический порог — профилирование 1/10/50 активных машин в SP и MP с измерением времени RoadcraftHooks, allocations/sec, GC pause и среднего/99-го перцентиля frame time. Только после этого имеет смысл переходить на MethodHandle или переиспользуемые mutable-снимки.",
    )


def add_tests(doc: Document, compatibility, bullet_id: int):
    doc.add_heading("10. Проверки и статус готовности", level=1)
    jar_hash = sha256(JAR) if JAR.exists() else "JAR отсутствовал при генерации документа"
    rows = [
        ["Сборка", "11 main sources; JDK 24, -Xlint:all,-output-file-clash -Werror", "ПРОШЛО"],
        ["Математическое ядро", "44 105 assertions: передачи, тяга, сцепление, тормоза, steering, NaN/границы", "ПРОШЛО"],
        ["Performance smoke", "1/10/50 машин, 10 000 кадров каждой группы", "ПРОШЛО; не game benchmark"],
        ["Lua/Java bridge contract", "24 assertions, включая fallback значений и fail-safe после трёх ошибок", "ПРОШЛО"],
        ["Transformer contract", "25 assertions по сигнатурам/callsite и fail-closed", "ПРОШЛО"],
        ["Реальный -javaagent smoke", "Game Java 25, загрузка точных B42.20.4 классов без запуска игры", "ACTIVE; ПРОШЛО"],
        ["SP в игре", "Управление, Bullet, столкновения, save/load", "НЕ ВЫПОЛНЕНО"],
        ["Listen MP", "Host + удалённый клиент, authority transfer", "НЕ ВЫПОЛНЕНО"],
        ["Dedicated MP", "Сервер + клиент, server launch agent", "НЕ ВЫПОЛНЕНО"],
    ]
    add_table(doc, ["Проверка", "Охват", "Результат"], rows, [2300, 4750, 2310], font_size=8.9)

    doc.add_heading("10.1. Идентификация проверенной сборки", level=2)
    identity = [
        ["Project Zomboid", f"{compatibility['knownTested']['gameVersion']} {compatibility['knownTested']['gameBuild']}"],
        ["projectzomboid.jar SHA-256", compatibility["knownTested"]["gameJarSha256"]],
        ["roadcraft-agent.jar SHA-256", jar_hash],
        ["Java runtime тестов", "Project Zomboid Java 25.0.1"],
        ["PZBullet", "1.0.0.28"],
    ]
    add_table(doc, ["Артефакт", "Значение"], identity, [2900, 6460], font_size=8.7)

    doc.add_heading("10.2. Минимальная игровая приёмка", level=2)
    for text in [
        "SP: три типа машин, холодный/запущенный двигатель, автомат/ручная КПП, задний ход и максимальная скорость.",
        "Покрытие: сухой асфальт, дождь, снег и грунт; целые/изношенные/снятые шины.",
        "Манёвры: экстренное торможение, ручник, burnout, резкий руль на 20/80/120 км/ч.",
        "Столкновения: растение, один зомби, группа, труп; сравнить при множителях 0/0,5/1/2.",
        "Буксировка: лёгкий и тяжёлый прицеп, подъём, торможение, разворот, смена водителя.",
        "Save/load и отключение мода: сохранение запускается, а INCOMPATIBLE/disabled оставляет ванильное поведение.",
        "MP: listen и dedicated, host/клиент, пересадка, reconnect и отсутствие двойной силы.",
    ]:
        add_list_item(doc, text, bullet_id)


def add_limitations(doc: Document, bullet_id: int):
    doc.add_heading("11. Ограничения, исключения и дальнейшая работа", level=1)
    risks = [
        ["Калибровка единиц", "Перевод рассчитанных ньютонов в engineForce/brakingForce использует новые калибровочные коэффициенты; ощущения ещё не сверены в игре.", "Высокий"],
        ["Reflection B42.20.4", "Адаптер кеширован, но привязан к текущим полям/методам; структурный transformer не доказывает, что все private-access пути сохранились в неизвестном патче.", "Высокий"],
        ["MP authority", "Проверка isLocalPhysicSim реализована, но переход authority при лаге/reconnect не испытан.", "Высокий"],
        ["Tire visuals", "skidInfo и состояние колёс требуют визуального/физического подтверждения.", "Средний"],
        ["Категоризация", "Sport/Standard/Heavy выводится из массы и мощности VehicleScript, без таблицы отдельных моделей.", "Средний"],
        ["Производительность", "Математика быстрая, но полный adapter/native/MP профиль отсутствует.", "Средний"],
        ["Сторонние моды", "Другие Java-agent/ASM-патчи тех же методов могут конфликтовать; порядок трансформаций не проверен.", "Средний"],
    ]
    add_table(doc, ["Риск", "Содержание", "Приоритет"], risks, [2100, 5600, 1660], font_size=9)

    doc.add_heading("11.1. Рекомендуемый порядок следующей работы", level=2)
    order = [
        "Запустить SP smoke на чистом профиле B42.20.4 и собрать полный console.txt.",
        "Исправить только подтверждённые runtime-ошибки доступа/единиц; повторить автоматические тесты.",
        "Настроить тягу, тормоза, steering и сопротивления по трём эталонным машинам.",
        "Провести профилирование 1/10/50 машин; оптимизировать hot path только по данным.",
        "Проверить listen MP, затем dedicated MP, включая перенос authority.",
        "После успешной матрицы сменить 0.1.0-dev на первую публичную prerelease-версию.",
    ]
    decimal_id = create_numbering(doc, "decimal", "%1.")
    for text in order:
        add_list_item(doc, text, decimal_id)


def add_provenance(doc: Document, bullet_id: int):
    doc.add_heading("12. Происхождение и лицензионная граница", level=1)
    add_body_paragraph(
        doc,
        "Новая кодовая база создана в отдельном workspace. Референс использовался для составления поведенческой спецификации и выявления опасной архитектуры полной подмены классов; его файлы не входят в поставку.",
    )
    for text in [
        "Ни один старый .class, .java, .lua, .wav, .png или текст установки не копируется в итоговый архив.",
        "В новом agent JAR допускаются только dev/roadcraft/pz/** и служебный META-INF; zombie/** отсутствует.",
        "Имена пакетов, продукта, ключей, локализации и исходный код новые.",
        "Сгенерированные декомпилированные материалы использовались только как временный локальный анализ API и удаляются перед выпуском.",
        "Исходники новой реализации предлагаются под MIT License; лицензия Project Zomboid и права правообладателей игры остаются отдельными.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_callout(
        doc,
        "Итог",
        "Текущий результат — технически собранный prerelease общей B42-ветки с безопасным отказом. Он готов к первому игровому smoke-тесту, но пока не должен называться подтверждённо рабочим SP/MP-релизом.",
        kind="positive",
    )


def audit_document(doc: Document):
    section = doc.sections[0]
    checks = {
        "page width": section.page_width == Inches(8.5),
        "page height": section.page_height == Inches(11),
        "top margin": section.top_margin == Inches(1),
        "right margin": section.right_margin == Inches(1),
        "bottom margin": section.bottom_margin == Inches(1),
        "left margin": section.left_margin == Inches(1),
        # Word stores these values in whole twips; 0.492 in rounds to the nearest twip.
        "header distance": abs(section.header_distance - Inches(0.492)) <= 635,
        "footer distance": abs(section.footer_distance - Inches(0.492)) <= 635,
    }
    failed = [name for name, ok in checks.items() if not ok]
    if failed:
        raise RuntimeError("Preset audit failed: " + ", ".join(failed))
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        if tbl_w is None or int(tbl_w.get(qn("w:w"), "0")) != CONTENT_WIDTH_DXA:
            raise RuntimeError("Table width audit failed")
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        if tbl_ind is None or int(tbl_ind.get(qn("w:w"), "-1")) != TABLE_INDENT_DXA:
            raise RuntimeError("Table indent audit failed")


def main():
    compatibility = read_compatibility()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.core_properties.title = "Roadcraft Dynamics — техническое описание B42"
    doc.core_properties.subject = "Независимая физика транспорта для Project Zomboid Build 42"
    doc.core_properties.author = "Project contributors"
    doc.core_properties.keywords = "Project Zomboid, Build 42, vehicle physics, Java agent, Roadcraft Dynamics"
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)

    bullet_id = create_numbering(doc, "bullet", "•")
    decimal_id = create_numbering(doc, "decimal", "%1.")

    add_cover(doc, compatibility)
    add_contents(doc, decimal_id)
    add_scope(doc, bullet_id)
    add_functional_scope(doc, bullet_id)
    add_architecture(doc, bullet_id)
    add_physics_details(doc, bullet_id)
    add_settings(doc)
    add_compatibility(doc, compatibility, bullet_id)
    add_installation(doc, decimal_id, bullet_id)
    add_multiplayer(doc, bullet_id)
    add_performance(doc, bullet_id)
    add_tests(doc, compatibility, bullet_id)
    add_limitations(doc, bullet_id)
    add_provenance(doc, bullet_id)

    audit_document(doc)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")
    print(f"Sandbox settings documented: {len(read_settings_rows())}")
    print(f"Agent SHA-256: {sha256(JAR)}")


if __name__ == "__main__":
    main()
